

######################### REAEDME.md #####################################

### license : Creative Commons Attribution-NonCommercial-NoDerivatives (CC BY-NC-ND)
### https://creativecommons.org/licenses/by-nc-nd/4.0/
### Made by Woo Yoon Kyu


#########################################################################
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
from langgraph.graph import StateGraph
from typing import TypedDict, Literal
import fitz  # PyMuPDF
import tempfile
from graph import *  # 요약 & 평가 로직
from dotenv import load_dotenv
import os
import re

load_dotenv()






# ====== STREAMLIT 설정 ======
st.set_page_config(page_title="멀티 논문 요약기", layout="wide")


# ✅ 스타일: 최대 너비 제한 (CSS 삽입)
# Streamlit 상단에 CSS 추가
st.markdown("""
<style>
    .main {
        max-width: 1000px;
        margin: 0 auto;
        padding-top: 2rem;
    }
    .element-container textarea, .element-container pre, .element-container code {
        white-space: pre-wrap !important;
        word-wrap: break-word !important;
        overflow-x: auto !important;
    }
    .stTextArea > div > textarea {
        font-family: '맑은 고딕', sans-serif;
        font-size: 0.92rem;
        line-height: 1.4;
    }
</style>
""", unsafe_allow_html=True)



#################


BASE_DIR = os.path.dirname(os.path.abspath(__file__))  # 현재 파일 기준 디렉토리
logo_path = os.path.join(BASE_DIR, "IIPL.PNG")
license_path = os.path.join(BASE_DIR, "license.png")


st.title("📚 멀티 논문 요약 및 과제 정렬 분석기")

with st.container():
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("🧠 논문 요약 멀티 에이전트")
        st.image(logo_path, width=1000)  # 로고 경로 수정 필요
    with col2:
        st.image(license_path, width=200)

st.markdown("""
- 여러 PDF 논문을 업로드하면, 각 논문에 대해:
    - 과제 정렬 분석 결과
    - 최종 3줄 요약
  을 화면에 보여줍니다.
- 1차 요약, 2차 설명 등 상세 내용은 Word로 다운로드하세요.
""")


# ======= md, WORD 저장 =======
def save_summary_to_md(text: str, title: str) -> str:
    md_content = f"# {title}\n\n" + text.strip()
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode="w", encoding="utf-8")
    temp_file.write(md_content)
    temp_file.close()
    return temp_file.name


def save_summary_to_word(summary: str, title: str) -> str:
    doc = Document()

    # 제목
    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run("\n" + "―" * 50)

    # 본문 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # ✅ 한글 깨짐 방지 설정
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    for line in summary.strip().split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("1)") or line.startswith("2)") or line.startswith("3)") or line.startswith("4)"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif line.startswith("(1)") or line.startswith("(2)") or line.startswith("(3)") or line.startswith("(4)"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            doc.add_paragraph(line)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# ======= 파일 업로드 =======
# 상태 초기화
if "download_data" not in st.session_state:
    st.session_state.download_data = []
if "results" not in st.session_state:
    st.session_state.results = {}

uploaded_files = st.file_uploader("📤 논문 PDF 업로드 (복수 가능)", type=["pdf"], accept_multiple_files=True)
project_goal = st.text_area("🎯 과제 목적 입력 (모든 논문에 공통 적용)", placeholder= "한국어 의료 데이터베이스 기반 검색 및 질의응답을 위한 모델 성능 향상", height=120)
project_status = st.text_area(
    "📅 과제의 진행 상황을 입력해주세요",
    help="아래 형식을 따르세요:\n- [완료된 단계 (~MM/DD)] : 이미 수행된 작업\n- [예정된 단계 (~MM/DD)] : 앞으로 수행할 계획인 작업",
    placeholder="""
📌 입력 형식 예시:
- [완료된 단계 (~MM/DD)] : 이미 수행한 작업
- [예정된 단계 (~MM/DD)] : 앞으로 수행할 계획인 작업

#### 진행 상황 ####

[데이터 수집 및 실험 완료 (~4/6)]  # 완료된 단계
<QA 시스템>
- Dataset: snuh/ClinicalQA
- 사용 모델: Exaone-7.8B-instruct, EEVE-10.8B, QWEN 2.5 7B, OpenBioLLM + chatvector + Blocking + LoRA
<Dialogue Generation>
- Dataset: 내부 커스텀 데이터셋
- 사용 모델: GPT-4o

[의료 도메인 데이터 추가 크롤링 (~4/13)]  # 예정된 단계
- 대한고혈압학회, 당뇨병학회, 치매 관련 협회 등에서 정보 수집
- 기존 QA 구조에 통합 가능하도록 전처리

[ko-LLM 대상 기법 적용 및 실험 (~4/27)]  # 예정된 단계
- 예정 기법:
- Chat Vector
- 특수문자 Unicode 블로킹
- CoT Steering
- Self-Consistency (또는 기타 파인튜닝 기법)
""",
    height=700
)

if uploaded_files and st.button("🚀 전체 요약 실행"):
    st.session_state.download_data = []
    st.session_state.results = {}

    for idx, uploaded_file in enumerate(uploaded_files, start=1):
        st.markdown(f"---\n## 📄 문서 {idx}: `{uploaded_file.name}`")
        progress = st.progress(0)
        status_text = st.empty()

        status_text.markdown("🔍 텍스트 추출 중...")
        title, body = extract_text_from_pdf(uploaded_file)
        progress.progress(10)

        status_text.markdown("📄 요약/비평 루프 진행 중...")
        graph = build_graph()
        state = {
            "document": body,
            "summary": "",
            "expanded_summary": "",
            "feedback": "",
            "critic_result": "",
            "status": "in_progress",
            "loop_count": 0,
            "project_goal": project_goal.strip(),
            "project_status": project_status.strip()
        }

        while state["status"] != "done" and state["loop_count"] < 5:
            state = graph.invoke(copy.deepcopy(state))
            progress.progress(min(90, 40 + state["loop_count"] * 10))
            status_text.markdown(f"🔁 {state['loop_count']}회차 평가 진행 중...")

        progress.progress(100)
        status_text.markdown("✅ 완료!")

        # 화면 출력용 데이터 저장
        st.session_state.results[title] = {
            "alignment": state["alignment_analysis"],
            "summary_pro": state["summary_pro_3line"],
            "summary_nonpro": state["summary_nonpro_3line"]
        }

        # 다운로드용 데이터 저장
        summary_doc = save_summary_to_word(state["summary"], f"{title} - 1차 요약")
        expanded_doc = save_summary_to_word(state["expanded_summary"], f"{title} - 2차 설명")
        align_doc = save_summary_to_word(state["alignment_analysis"], f"{title} - 과제 정렬 분석")

        # ✅ 3줄 요약 통합 저장 (Word/MD)
        summary_3line_text = f"[전문가용 요약]\n{state['summary_pro_3line']}\n\n[일반인용 요약]\n{state['summary_nonpro_3line']}"
        summary_3line_docx = save_summary_to_word(summary_3line_text, f"{title} - 최종 3줄 요약")
        summary_3line_md = save_summary_to_md(summary_3line_text, f"{title} - 최종 3줄 요약")

        st.session_state.download_data.append({
            "title": title,
            "summary": summary_doc,
            "expanded": expanded_doc,
            "alignment": align_doc,
            "summary_3line_docx": summary_3line_docx,
            "summary_3line_md": summary_3line_md
        })

# ✅ 결과 화면 출력
if st.session_state.results:
    st.markdown("---")
    
    align_dict = {title: res["alignment"] for title, res in st.session_state.results.items()}
    overview_result = call_alignment_overview_agent(
        project_goal=project_goal,
        project_status=project_status,
        alignments=align_dict
    )


    st.markdown("## 📋 전체 문서 종합 정렬 평가 요약")
    st.markdown(overview_result)
    st.markdown("---")
    
    st.markdown("## 📊 세부 분석 결과 보기")
    for i, (title, res) in enumerate(st.session_state.results.items(), start=1):
        st.markdown(f"### 📄 문서 {i}: `{title}`")
        st.subheader("🔗 과제 정렬 분석 결과")
        st.markdown(res["alignment"])
        st.subheader("🧾 최종 3줄 요약")
        st.markdown("**전문가용 요약**")
        st.markdown(res["summary_pro"])
        st.markdown("**일반인용 요약**")
        st.markdown(res["summary_nonpro"])

# ✅ 다운로드 영역 유지
if st.session_state.download_data:
    st.markdown("---")
    st.markdown("## 📥 모든 문서 요약 다운로드")
    for i, item in enumerate(st.session_state.download_data, start=1):
        st.markdown(f"### 📄 문서 {i}: `{item['title']}`")
        col1, col2, col3 = st.columns(3)
        with col1:
            with open(item["summary"], "rb") as f:
                st.download_button("📝 1차 요약 다운로드", f, file_name=f"{item['title']}_1차요약.docx")
        with col2:
            with open(item["expanded"], "rb") as f:
                st.download_button("💬 2차 설명 다운로드", f, file_name=f"{item['title']}_2차설명.docx")
        with col3:
            with open(item["alignment"], "rb") as f:
                st.download_button("🔗 정렬 분석 다운로드", f, file_name=f"{item['title']}_정렬분석.docx")

        col4, col5 = st.columns(2)
        with col4:
            with open(item["summary_3line_docx"], "rb") as f:
                st.download_button("📌 3줄 요약 다운로드 (.docx)", f, file_name=f"{item['title']}_3줄요약.docx")
        with col5:
            with open(item["summary_3line_md"], "rb") as f:
                st.download_button("📄 3줄 요약 다운로드 (.md)", f, file_name=f"{item['title']}_3줄요약.md")