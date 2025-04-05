# %%
import streamlit as st
from langgraph.graph import StateGraph
from typing import TypedDict, Literal
import fitz  # PyMuPDF
from docx import Document
import tempfile
from openai import OpenAI
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing_extensions import Annotated  # Python 3.9+ or install via pip

# ======= PROMPTS =======
summarizer_prompt = """
당신은 AI 전문가입니다. 다음 문서(또는 논문)를 기반으로, 다른 AI 전문가가 핵심을 빠르게 파악할 수 있도록 **구조화된 요약**을 작성해주세요.

아래 4개의 항목에 대해 반드시 작성하되, 각 항목은 **Subsection 제목 + Bullet 형식**으로 구성해주세요.  
특히 **motivation, research gap, contributions, methodology overview, step-by-step method, experiment summary** 등을 구체적으로 구분해주세요.  
각 **Subsetsion**에는 최소 5 문장의 내용이 들어가야 합니다.

---

(1) **개요 (Introduction / Background)**  
- 핵심 Subsections:
    - 📌 *Motivation* (왜 이 연구가 필요한가?)  
    - 📌 *Research Gap* (기존 연구와의 차별성은?)  
    - 📌 *Contributions* (논문의 주요 기여는?)  
    - 📌 *Related Works* (참고 문헌과의 직접적인 연결 중심 — 어떤 연구의 어떤 내용을 계승/차별화했는가?)  
- 각 항목은 **Bullet List**로 작성  
- 기술 용어는 반드시 (괄호 안에 간단히 해설) 포함

(2) **방법론 (Methodology)**  
- 핵심 Subsections:
    - ⚙️ *Overview*: 전체 구조 요약  
    - ⚙️ *Components*: 주요 구성 요소 및 알고리즘  
    - ⚙️ *Step-by-Step*: 방법론의 실행 단계  
    - ⚙️ *Differences*: 기존 기법 대비 구조적 차이  
- 가능한 경우 수식이나 아키텍처를 언어로 설명  
- **각 Subsection은 Bullet 형식으로 정리**

(3) **장점 및 효과 (Advantages and Impacts)**  
- 핵심 Subsections:
    - ✅ *Technical Benefits* (정확도, 속도, 확장성 등)  
    - ✅ *Practical Use Cases* (실제 산업 적용 가능성을 중심으로 — 현재 기술 트렌드나 산업 분야와 연결)  
    - ✅ *Expected Impact* (과제나 업무 수행 시 기대 효과, 산업적 또는 사회적 가치 포함)  
    - ✅ *Evaluation Summary* (사용한 주요 metric과 그 결과 — 예: 평균 응답 속도 20% 향상, RMSE 15% 감소 등 **숫자 포함 간결한 요약**)  
    - ✅ *Limitations* (알려진 한계나 고려사항)  
- 각 항목은 Bullet로 짧고 명확하게 정리

(4) **예상 질문 및 정답 (Q&A / FAQ)**  
- 기술을 처음 접하는 전문가들이 가질 법한 질문을 3~5개 상정  
- 각 질문은 ❓ 로 시작, 답변은 ➤ 로 시작  
- 예시:
    - ❓ 이 방법이 기존에 진행되었던 연구들보다 정확한 이유는?
    - ➤ 기존 연구들은 [...]이 부족했지만, 본 연구는 [...]를 개선함.

---

**작성 지침**  
- 각 항목은 반드시 (1), (2), (3), (4) 형식으로 구분  
- 모든 기술 용어는 (괄호 안에 간단히 해설) 추가  
- 전체 길이는 18000자 이내  
- 참고 문헌은 관련 논문을 포함하여 **APA 형식**으로 맨 아래 정리  
- 참고 문헌은 **가능하면 7개 이상**, 다양한 분야에서 관련 연구를 폭넓게 제시해줘.  
- 이 결과는 “일반인용 설명” 생성을 위한 입력으로 사용됩니다.
"""


expander_prompt = """
다음은 한 논문의 원문, 해당 논문을 요약한 전문가용 해설, 그리고 이 내용을 참고해야 할 과제의 목적이야.

당신의 역할은 기술을 쉽게 풀어 설명하는 것뿐 아니라,  
**"이 기술이 실제로 해당 과제를 해결하는 데 실질적인 도움이 될 수 있는가?"를 먼저 판단**하고,  
그 판단에 근거해 일반인이 이해하기 쉬운 설명을 작성하는 것이야.

따라서 설명은 단순히 내용을 쉽게 풀어내는 것을 넘어서,  
**"왜 이 기술이 이 과제에 의미 있고 쓸모 있는지"를 독자가 자연스럽게 이해할 수 있도록 구성**해야 해.

---

(1) **개요**  
   - 어려운 개념이나 용어는 풀어서 설명하고, 가능한 한 비유나 사례를 활용해줘.  
   - 배경보다는 “이게 왜 중요한지”와 “무엇을 가능하게 해주는지”에 더 집중해줘.  
   - 특히, **이 기술이 과제 목적과 실제로 어떤 방식으로 연결되는지를 설명해줘**.  
     연결이 약하다면 그 이유도 함께 언급하고, 가능한 대안적 기술도 제안해줘.  
   - 기술적인 핵심도 빠뜨리지 말고 간단하게라도 구조적 방식(예: "데이터를 이렇게 처리하고, 이렇게 결합한다" 식)을 넣어줘.  
   - 특히 "방법론"의 핵심 아이디어가 빠지지 않도록 해줘. 그래야 독자가 '이건 기술적으로 새로운 방식이구나'라는 느낌을 받을 수 있어.

(2) **장점 및 효과**  
   - 이 기술이 어떤 문제를 어떻게 해결하는지, 그리고 그게 왜 유용한지를 실생활이나 과제/업무 관점에서 풀어줘.  
   - 사람들이 실제로 **이 기술을 어디에 쓸 수 있는지, 어떤 결과를 기대할 수 있는지**에 초점을 맞춰 설명해줘.  
   - 특히, **과제에서 해결하려는 문제와 이 기술이 어떻게 연결되는지를 명확히 표현해줘.**  
     직접적인 기여가 어려울 경우, 간접적인 연관 가능성이나 고려할 점도 함께 언급해줘.  
   - 설명이 너무 일반적이지 않도록, 기술적인 핵심을 간단하게라도 설명해줘 (예: "데이터를 이렇게 처리하고, 이렇게 결합한다").  
   - "방법론"이 단지 존재한다고만 하지 말고, 어떤 방식으로 다르고 왜 그게 효과적인지도 표현해줘.

(3) **예상 질문 및 정답**  
   - 기술에 대해 일반인이 가질 수 있는 현실적인 의문을 상정하고, 최대한 쉽게, 구체적으로 답변해줘.  
   - 기술적인 설명보다는 “이게 나한테 어떤 의미가 있지?”를 해소해주는 식이면 좋아.  
   - 과제 수행 중 생길 수 있는 **기술 적용의 현실적인 걱정이나 불확실성**도 질문으로 포함해줘.  
   - 앞서 설명한 장점과 중복되지 않도록, 질문은 주로 기술 적용 시의 현실적 의문 위주로 작성해줘.

**형식 조건**  
- 항목 구조는 반드시 (1), (2), (3) 형식을 포함하고, 유지할 것  
- 전체 길이는 10000자 이내  
- 표현은 쉽고 친근하게, 핵심은 빠짐없이  
"""




critic_prompt = """
다음은 논문의 원문, 전문가 요약, 일반인을 위한 2차 쉬운 설명, 그리고 이 내용을 참고해 작성된 과제의 목적이야.

당신의 역할은 AI 전문가로서 다음 두 문서를 평가해줄 사람이야:

1. 전문가를 위한 요약 (Summarizer)
2. 일반인을 위한 설명 (Expander)

특히 이 문서들이 단순히 내용을 정리한 것을 넘어서,  
**"이 기술이 실제로 이 과제를 해결하는 데 도움이 될 수 있는가?"**  
라는 관점에서 평가하는 것이 중요해.


---  

1) **전문가 요약 평가 (Summarizer Evaluation)**  
- (1) 문서의 핵심 내용을 충실히 요약했는가? (O/X)  
- (2) 형식 요구사항(1, 2, 3, 4 구조 및 분량 등)을 잘 지켰는가? (O/X)  
- (3) 참고 문헌이 실제 존재하거나 논문과 관련성이 높은가? (O/X)  
- (4) 과제 목적을 기준으로 보았을 때, 요약된 기술이 **실제로 과제 해결에 사용될 수 있거나, 적용 가능성에 대한 구체적 연결이 존재하는가?** (O/X)  

→ 이 항목은 단순히 “언급”이 아닌, **기술이 과제에 어떤 방식으로 기여할 수 있는지를 납득할 수 있어야** O를 줄 수 있어.

2) **일반인용 설명 평가 (Expander Evaluation)**  
- (1) 전문가 요약의 핵심 내용을 잘 반영했는가? (O/X)  
- (2) 표현이 일반인이 이해할 수 있을 정도로 충분히 쉬운가? (O/X)  
- (3) 중요한 기술 포인트나 효과가 빠지지 않았는가? (O/X)  
- (4) 이 설명을 읽었을 때, 일반인도 **“이 기술이 과제를 해결하는 데 실질적인 도움이 되겠구나”라고 느낄 수 있는가?** (O/X)  

→ 기술적 설명보다 중요한 건, **과제 목적에 비춰 봤을 때 이 기술이 실제로 쓸모 있어 보이느냐**는 점이야.

3) **과제 정렬 분석 평가 (Project Alignment Evaluation)**
- (1) 과제의 기술적 흐름을 잘 요약했는가? (O/X)
- (2) 논문 방법론이 기여할 수 있는 지점을 실제로 제시했는가? (O/X)
- (3) 효과와 한계를 균형 있게 설명했는가? (O/X)

→ 특히 “진행 상황”과 “논문 기술” 간의 연결성이 납득될 정도여야 함

---

각 항목에 대해 O/X로 평가하고,  
**문제가 있는 경우는 반드시 그 이유를 구체적으로 설명**해주세요.  
(예: 연관성 부족, 논리 비약, 설명 불충분, 적용 가능성 모호함 등)

**마지막 줄에는 반드시 다음 중 하나만 작성해주세요:**  
[완료] 또는 [수정 필요]
"""


project_aligner_prompt = """
다음은 한 논문 원문, 전문가용 요약, 그리고 현재 과제의 목적과 기술적 진행 상황이다.

당신은 AI 기술 적용 컨설턴트다. 지금의 과제 진행 흐름과 실험 방향성을 고려했을 때,  
**이 논문이 제안하는 방법론이 과제에서 어떤 문제를 해결하거나 보완할 수 있는지를 분석**해줘.

입력된 '과제 진행 상황'은 다음과 같은 형식을 따르고 있어:
- [완료된 단계 (~MM/DD)] 형식은 이미 수행한 내용을 의미함
- [예정된 단계 (~MM/DD)] 형식은 앞으로 수행할 계획임
- 따라서 각 단계를 구분해 분석하고, 논문 방법론이 '예정된 단계'에 어떤 도움이 될 수 있는지 특히 중점적으로 언급해줘.
- 예정된 단계는 분석의 중심이 되며, 이 단계에서 논문 기술이 **직접적으로 기여할 수 있는 부분을 구체적으로 언급**해야 해.

📌 단, 각 예정된 단계에 대해 논문 기술이 **직접적 또는 간접적으로 기여할 수 있는지 개별적으로 판단**해줘.  
기여 가능성이 낮거나 연관이 적은 단계는 **분석에서 제외하고**, 그 이유도 간단히 설명해줘.  
억지로 연결하지 말고, **정말 기여가 예상되는 부분에만 집중**해줘.

🛑 논문이 과제의 흐름과 맞지 않거나, 현재 예정 단계에서 **기여할 수 있는 부분이 없다고 판단되면**,  
그 이유를 구체적으로 설명하고, **어떤 조건이나 시점에서 다시 검토하면 좋을지도 제안**해줘.

- 단순한 요약이 아니라, “어떤 실험 단계나 기술 흐름에서 이 방법론이 유의미하게 작용할 수 있을지”를 구체적으로 제시해줘.
- 현재 사용 중인 모델, 데이터 흐름, 적용하려는 기술들과 논문 속 방법론이 어떻게 연결될 수 있는지도 언급해줘.

출력 형식은 다음과 같아:
1. 과제 기술 흐름 요약
2. 논문 방법론이 기여 가능한 구체적 지점
3. 예상되는 적용 효과
4. 한계나 고려사항
"""

# ✅ 전체 정렬 평가용 프롬프트 (표 생성 포함)
alignment_overview_prompt = """
당신은 AI 과제 설계 및 기술 적용 전략가입니다.

아래는 하나의 프로젝트 목적과 진행 상황, 그리고 여러 논문에 대해 분석된 과제 정렬 평가 내용입니다.

각 논문이 이 과제의 현재 흐름과 얼마나 잘 맞는지, 다음 정보를 표 형식으로 제공해주세요:
1. 적용 가능성 (높음/중간/낮음)
2. 한 줄 요약 (이 논문이 어떻게 기여할 수 있는지)
3. 활용 추천 시점 (예: 데이터 수집 단계, 모델 학습 단계, QA 응답 생성 단계 등)

표 형식 예시:

| 논문 제목 | 적용 가능성 | 한 줄 요약 | 활용 추천 시점 |
|-----------|--------------|-----------------------------|--------------------|
| 논문 A    | 높음 ✅     | 실험 구조와 정확도 향상에 직접 기여 가능 | 모델 학습 단계 |
| 논문 B    | 낮음 ❌     | 기술 흐름 차이로 인해 직접 적용 어려움 | 활용 불가 |

- 표 형식은 Markdown 테이블로 출력
- 제목은 실제 논문 파일명 또는 문서 제목을 활용
- 표현은 명확하고 간결하게
- 논문이 쓸모 없다면 이유와 함께 '활용 불가'로 명시
"""

# chatgpt-4o-latest
# gpt-4o
# ======= AGENTS =======
def call_openai_gpt(prompt: str, model="chatgpt-4o-latest") -> str:
    client = OpenAI()
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "답변은 표준적이고 명확해야 하며, 일반인이 이해할 수 있도록 구성되어야 합니다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.0,
    )
    return response.choices[0].message.content

def call_summarizer(document: str, feedback: str = "") -> str:
    prompt = summarizer_prompt
    if feedback:
        prompt += f"\n\n[Critic Feedback:\n{feedback}]\n"
    return call_openai_gpt(prompt + "\n\n" + document)

def call_expander(summary: str, document: str, project_goal: str, project_status: str) -> str:
    prompt = expander_prompt + f"""

===== 과제 목적 =====
{project_goal}

===== 과제 진행 상황 =====
{project_status if project_status.strip() else "진행 상황 정보 없음"}

===== 원문 =====
{document}

===== 1차 요약 =====
{summary}
"""
    return call_openai_gpt(prompt)

def call_critic(document: str, summary: str, expanded: str, project_goal: str, alignment: str):
    prompt = critic_prompt + f"""

===== 과제 목적 =====
{project_goal}

===== 논문 원문 =====
{document}

===== 1차 요약 =====
{summary}

===== 2차 설명 =====
{expanded}

===== 과제 정렬 분석 =====
{alignment}
"""
    result = call_openai_gpt(prompt)
    ...

    if "[완료]" in result:
        return result, "done", ""
    else:
        return result, "in_progress", result

        # "summary_pro_3line": summary_pro,
        # "summary_nonpro_3line": summary_nonpro,


def call_project_aligner(summary: str, document: str, project_goal: str, project_status: str) -> str:
    prompt = project_aligner_prompt + f"""

===== 과제 목적 =====
{project_goal}

===== 과제 진행 상황 =====
{project_status if project_status.strip() else "진행 상황 정보 없음"}

===== 논문 요약 =====
{summary}

===== 논문 원문 =====
{document}
"""
    return call_openai_gpt(prompt)


# ✅ Agent 함수 정의 (전체 종합 분석)
def call_alignment_overview_agent(project_goal: str, project_status: str, alignments: dict[str, str]) -> str:
    prompt = alignment_overview_prompt + f"""

===== 과제 목적 =====
{project_goal}

===== 과제 진행 상황 =====
{project_status}

===== 논문별 정렬 평가 =====
""" + "\n\n".join([f"[{k}]\n{v}" for k, v in alignments.items()])

    return call_openai_gpt(prompt)


# ======= GRAPH =======
class State(TypedDict):
    document: Annotated[str, "static"]
    summary: str
    expanded_summary: str
    feedback: str
    status: Literal["in_progress", "done"]
    loop_count: int
    critic_result: str
    project_goal: str
    project_status: str  # ✅ 추가
    summary_pro_3line : str
    summary_nonpro_3line : str
    alignment_analysis: str  # 🔥 새 필드

def summarizer_node(state):
    summary = call_summarizer(state["document"], state.get("feedback", ""))
    return {**state, "summary": summary}

def expander_node(state):
    expanded = call_expander(
        state["summary"],
        state["document"],
        state["project_goal"],
        state.get("project_status", "")
    )
    return {**state, "expanded_summary": expanded}

def project_aligner_node(state):
    result = call_project_aligner(
        state["summary"],
        state["document"],
        state["project_goal"],
        state["project_status"]
    )
    return {**state, "alignment_analysis": result}


def critic_node(state):
    result, status, feedback = call_critic(
        state["document"],
        state["summary"],
        state["expanded_summary"],
        state["project_goal"],
        state.get("alignment_analysis", "")  # ✅ 이 줄이 꼭 필요해요
    )
    return {
        **state,
        "critic_result": result,
        "feedback": feedback,
        "status": status,
        "loop_count": state["loop_count"] + 1
    }


def check_continue(state):
    if state["status"] == "done":
        return "final"  # 3줄 요약으로 이동
    elif state["loop_count"] < 5:
        return "continue"  # 다시 summarizer로 루프
    else:
        return "exit"  # 최대 횟수 도달, 종료


def final_3line_summary_node(state):
    summary_pro = call_openai_gpt(
        f"""
        당신은 AI 논문 요약 전문가입니다.

        아래는 어떤 논문을 정리한 전문가용 요약입니다.  
        이 내용을 기반으로, 다음 3가지 항목에 대해 한 줄씩 요약해주세요.

        1. 문제 정의와 제안된 해결 방법  
        2. 제안된 방법론의 핵심 구조  
        3. 주요 실험 결과 또는 기대 효과  

        각 항목은 최소 세 항으로 띄어서 앞에 '-' 을 붙여서 보기 좋게 정리해주세요.

        ===== 전문가 요약 =====
        {state["summary"]}
        """
    )

    summary_nonpro = call_openai_gpt(
        f"""
        당신은 기술 내용을 일반인도 쉽게 이해할 수 있도록 설명하는 AI 전문가입니다.

        아래는 어떤 기술에 대한 일반인용 쉬운 설명과 과제 목적입니다.  
        이 내용을 바탕으로, 다음 3가지 항목을 각각 한 줄씩 설명해주세요:

        1. 이 기술이 어떤 방식인지 (쉽게 핵심 요약)  
        2. 이 기술이 과제 목적과 얼마나 잘 맞는지  
        3. 이 기술을 적용했을 때 기대되는 효과  

        각 항목은 비전공자가 이해할 수 있도록 간결하고 친근하게 설명해주세요.

        ===== 과제 목적 =====
        {state["project_goal"]}

        ===== 일반인용 설명 =====
        {state["expanded_summary"]}
        """
    )
    return {
        **state,
        "summary_pro_3line": summary_pro,
        "summary_nonpro_3line": summary_nonpro,
    }






def build_graph():
    builder = StateGraph(State)

    # 주요 노드 등록
    builder.add_node("summarizer", summarizer_node)
    builder.add_node("expander", expander_node)
    builder.add_node("critic", critic_node)
    builder.add_node("final_3line_summary", final_3line_summary_node)  # ✅ 추가
    builder.add_node("project_aligner", project_aligner_node)
    
    # 시작 지점
    builder.set_entry_point("summarizer")

    # 기본 흐름
    builder.add_edge("summarizer", "expander")
    # 흐름 순서 변경
    builder.add_edge("expander", "project_aligner")
    builder.add_edge("project_aligner", "critic")
    builder.add_edge("critic", "final_3line_summary")

    # critic 조건 분기 → 완료 시 final 요약 생성
    builder.add_conditional_edges("critic", check_continue, {
        "final": "final_3line_summary",
        "continue" : "summarizer",
        "exit": "__end__"  # ✅ 종료 시 반드시 "__end__" 사용
    })


    return builder.compile()


def extract_text_from_pdf(file) -> tuple[str, str]:
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    first_newline = text.find("\n")
    title = text[:first_newline].strip() if first_newline != -1 else "제목 없음"
    body = text[first_newline:].strip() if first_newline != -1 else text
    return title, body

def save_summary_to_word(summary: str, title: str) -> str:
    doc = Document()

    # 제목 스타일
    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run("\n" + "―" * 50)

    # 본문 스타일
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    for line in summary.strip().split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("1)") or line.startswith("2)") or line.startswith("3)"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            doc.add_paragraph(line)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name
# %%

